"""
ingestion.py
------------
Handles discovery, loading, parsing, and chunking of documents inside the
Knowledge Base directory. Supports PDF, DOCX, TXT, Markdown, and CSV files.

Each supported file is converted into one or more Chunk objects, which carry
the extracted text plus metadata (source file, chunk index) that is later
used for grounding and citation when the assistant answers a question.
"""

from __future__ import annotations

import csv
import hashlib
import os
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".markdown", ".csv"}


@dataclass
class Chunk:
    """A single unit of text extracted from a source document."""

    text: str
    source: str          # relative file path, e.g. "policies/leave_policy.pdf"
    chunk_index: int      # position of this chunk within the source file
    metadata: dict = field(default_factory=dict)

    @property
    def chunk_id(self) -> str:
        return f"{self.source}::chunk_{self.chunk_index}"


def file_hash(path: Path) -> str:
    """Return a hash of a file's contents + mtime, used to detect changes."""
    stat = path.stat()
    h = hashlib.sha256()
    h.update(str(stat.st_mtime_ns).encode())
    h.update(str(stat.st_size).encode())
    h.update(str(path).encode())
    return h.hexdigest()


def discover_files(kb_dir: Path) -> List[Path]:
    """Return all supported files inside the knowledge base directory (recursive)."""
    files = []
    for root, _dirs, filenames in os.walk(kb_dir):
        for name in filenames:
            path = Path(root) / name
            if path.suffix.lower() in SUPPORTED_EXTENSIONS:
                files.append(path)
    return sorted(files)


def _read_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_md(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _read_csv(path: Path) -> str:
    """Flatten a CSV file into readable text, one line per row with headers."""
    lines = []
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as f:
        reader = csv.reader(f)
        rows = list(reader)
    if not rows:
        return ""
    header = rows[0]
    for row in rows[1:]:
        pairs = [f"{h.strip()}: {v.strip()}" for h, v in zip(header, row)]
        lines.append(", ".join(pairs))
    return "\n".join(lines)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required to read PDF files. Install it with `pip install pypdf`."
        ) from exc

    reader = PdfReader(str(path))
    pages_text = []
    for page in reader.pages:
        pages_text.append(page.extract_text() or "")
    return "\n".join(pages_text)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to read DOCX files. Install it with `pip install python-docx`."
        ) from exc

    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


_READERS = {
    ".txt": _read_txt,
    ".md": _read_md,
    ".markdown": _read_md,
    ".csv": _read_csv,
    ".pdf": _read_pdf,
    ".docx": _read_docx,
}


def load_file_text(path: Path) -> str:
    """Dispatch to the correct reader based on file extension."""
    reader = _READERS.get(path.suffix.lower())
    if reader is None:
        raise ValueError(f"Unsupported file type: {path.suffix}")
    return reader(path)


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 150) -> List[str]:
    """
    Split text into overlapping chunks of approximately `chunk_size` characters.
    Splitting happens on paragraph boundaries where possible to keep chunks
    semantically coherent, falling back to hard splits for long paragraphs.
    """
    text = text.strip()
    if not text:
        return []

    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks: List[str] = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) + 1 <= chunk_size:
            current = f"{current}\n{para}" if current else para
        else:
            if current:
                chunks.append(current)
            if len(para) > chunk_size:
                # Hard-split very long paragraphs
                for i in range(0, len(para), chunk_size - overlap):
                    chunks.append(para[i:i + chunk_size])
                current = ""
            else:
                current = para
    if current:
        chunks.append(current)

    # Add overlap between consecutive chunks for better retrieval continuity
    overlapped = []
    for i, c in enumerate(chunks):
        if i == 0:
            overlapped.append(c)
        else:
            prev_tail = chunks[i - 1][-overlap:]
            overlapped.append(f"{prev_tail}\n{c}")
    return overlapped if overlapped else chunks


def build_chunks_for_file(path: Path, kb_dir: Path, chunk_size: int = 1000, overlap: int = 150) -> List[Chunk]:
    """Load a single file and split it into Chunk objects."""
    text = load_file_text(path)
    pieces = chunk_text(text, chunk_size=chunk_size, overlap=overlap)
    rel_source = str(path.relative_to(kb_dir))
    return [
        Chunk(text=piece, source=rel_source, chunk_index=i)
        for i, piece in enumerate(pieces)
    ]


def build_all_chunks(kb_dir: Path, chunk_size: int = 1000, overlap: int = 150) -> List[Chunk]:
    """Discover every supported file in kb_dir and return all resulting chunks."""
    all_chunks: List[Chunk] = []
    for file_path in discover_files(kb_dir):
        try:
            all_chunks.extend(
                build_chunks_for_file(file_path, kb_dir, chunk_size=chunk_size, overlap=overlap)
            )
        except Exception as exc:  # noqa: BLE001 - surfaced to the user, not swallowed
            print(f"[ingestion] Warning: failed to process '{file_path}': {exc}")
    return all_chunks
